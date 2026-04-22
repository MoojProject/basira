import io
import re
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter(prefix="/draft", tags=["Draft"])


class DraftRequest(BaseModel):
    content: str          
    title: str           
    contract_type: str   


def _build_docx(content: str, title: str) -> bytes:
    """
    يحوّل النص العادي إلى ملف Word منسّق:
    - العناوين (## أو السطور المنتهية بـ :) → Heading 2
    - الفقرات العادية → Normal
    - اتجاه RTL لكل فقرة
    """
    doc = Document()

    # ── إعدادات الصفحة (A4) ──
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.right_margin = Cm(3)      
    section.left_margin  = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    _set_rtl(h)

    doc.add_paragraph()  

    lines = content.strip().split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("##"):
            text = re.sub(r"^#+\s*", "", line)
            p = doc.add_heading(text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(13)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            _set_rtl(p)

        elif re.match(r"^(المادة|البند|الفصل|أولاً|ثانياً|ثالثاً|رابعاً|خامساً)\b", line):
            p = doc.add_heading(line, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(13)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            _set_rtl(p)

        elif line.endswith(":") and len(line) < 60:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2E, 0x4D, 0x6B)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(p)

        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(12)
            _set_rtl(p)

    doc.add_paragraph()
    doc.add_paragraph()
    sig_para = doc.add_paragraph("الطرف الأول: ________________          الطرف الثاني: ________________")
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sig_para.runs:
        sig_para.runs[0].font.name = "Arial"
        sig_para.runs[0].font.size = Pt(11)
    _set_rtl(sig_para)

    date_para = doc.add_paragraph("التاريخ: ________________")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date_para.runs:
        date_para.runs[0].font.name = "Arial"
        date_para.runs[0].font.size = Pt(11)
    _set_rtl(date_para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_rtl(paragraph):
    """يضع اتجاه RTL على الفقرة عبر XML مباشرة."""
    from docx.oxml.ns import qn
    from lxml import etree
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = etree.SubElement(pPr, qn("w:bidi"))
    bidi.set(qn("w:val"), "1")


@router.post("/generate-docx")
async def generate_docx(req: DraftRequest):
    """
    يستقبل النص المُولَّد ويرجع ملف .docx للتحميل المباشر.
    
    Body:
        content      : نص العقد الكامل
        title        : عنوان العقد (مثال: عقد عمل)
        contract_type: labor | rent | nda
    
    Response:
        application/vnd.openxmlformats-officedocument.wordprocessingml.document
    """
    docx_bytes = _build_docx(req.content, req.title)

    filename = f"basira_{req.contract_type}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )
